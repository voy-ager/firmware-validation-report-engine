import os
import uuid
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io
from .models import EnrichedMetrics

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "../templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "../reports")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def render(
    metrics: EnrichedMetrics,
    summary: dict,
    charts: dict,
    report_meta: dict,
) -> dict:
    context = _build_context(metrics, summary, charts, report_meta)
    pdf_path = _render_pdf(context)
    docx_path = _render_docx(context, metrics, charts)
    return {"pdf_path": pdf_path, "docx_path": docx_path}


def _build_context(metrics, summary, charts, meta):
    parsed = metrics.parsed
    return {
        "platform": parsed.platform,
        "build_id": parsed.build_id,
        "report_type": meta.get("report_type", "weekly"),
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "total_tests": parsed.total_tests,
        "total_passed": parsed.total_passed,
        "total_failed": parsed.total_failed,
        "total_skipped": parsed.total_skipped,
        "pass_rate": metrics.pass_rate,
        "fail_rate": metrics.fail_rate,
        "risk_score": metrics.risk_score,
        "risk_color": metrics.risk_color,
        "top_failures": metrics.top_failures,
        "failure_by_category": metrics.failure_by_category,
        "executive_summary": summary.get("executive_summary", ""),
        "risk_assessment": summary.get("risk_assessment", ""),
        "summary_ai_generated": summary.get("ai_generated", True),
        "risk_ai_generated": summary.get("ai_generated", True),
        "engineer_notes": meta.get("engineer_notes", ""),
        "approved_by": meta.get("approved_by", ""),
        "approved_at": meta.get("approved_at", ""),
        "charts": charts,
    }


def _render_pdf(context: dict) -> str:
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))
    template = env.get_template("report.html")
    html_content = template.render(**context)

    filename = f"report_{context['build_id']}_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(OUTPUT_DIR, filename)

    with open(output_path, "wb") as pdf_file:
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_file)

    if pisa_status.err:
        raise RuntimeError(f"PDF generation failed: {pisa_status.err}")

    return output_path


def _render_docx(context: dict, metrics: EnrichedMetrics, charts: dict) -> str:
    doc = Document()

    # Title
    title = doc.add_heading("Firmware Validation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.style = "Table Grid"
    headers = ["Platform", "Build ID", "Report Type", "Generated"]
    values = [
        context["platform"],
        context["build_id"],
        context["report_type"].upper(),
        context["generated_at"],
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        meta_table.cell(0, i).text = h
        meta_table.cell(1, i).text = v
    doc.add_paragraph()

    # Risk banner
    risk_para = doc.add_paragraph()
    risk_run = risk_para.add_run(
        f"RISK: {context['risk_score'].upper()} — "
        f"Pass Rate: {context['pass_rate']}% | "
        f"Failed: {context['total_failed']}/{context['total_tests']}"
    )
    risk_run.bold = True
    risk_colors = {"low": RGBColor(0x16, 0xa3, 0x4a), "medium": RGBColor(0xd9, 0x77, 0x06),
                   "high": RGBColor(0xea, 0x58, 0x0c), "critical": RGBColor(0xdc, 0x26, 0x26)}
    risk_run.font.color.rgb = risk_colors.get(context["risk_score"], RGBColor(0, 0, 0))
    risk_run.font.size = Pt(13)

    # Metrics
    doc.add_heading("Test Metrics", level=1)
    metrics_table = doc.add_table(rows=2, cols=4)
    metrics_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Total Tests", context["total_tests"]),
        ("Passed", context["total_passed"]),
        ("Failed", context["total_failed"]),
        ("Pass Rate", f"{context['pass_rate']}%"),
    ]):
        metrics_table.cell(0, i).text = label
        metrics_table.cell(1, i).text = str(val)
    doc.add_paragraph()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    ai_note = " [AI Generated]" if context["summary_ai_generated"] else ""
    doc.add_paragraph(context["executive_summary"] + ai_note)

    # Risk assessment
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph(context["risk_assessment"])

    # Engineer notes
    if context.get("engineer_notes"):
        doc.add_heading("Engineer Notes", level=1)
        doc.add_paragraph(context["engineer_notes"])

    # Charts
    doc.add_heading("Test Results — Charts", level=1)
    for chart_name, chart_b64 in charts.items():
        img_bytes = base64.b64decode(chart_b64)
        img_stream = io.BytesIO(img_bytes)
        doc.add_picture(img_stream, width=Inches(5.5))
        doc.add_paragraph()

    # Failures table
    if context["top_failures"]:
        doc.add_heading("Top Failures", level=1)
        fail_table = doc.add_table(rows=1, cols=4)
        fail_table.style = "Table Grid"
        headers = ["Test Name", "Category", "Status", "Failure Message"]
        for i, h in enumerate(headers):
            fail_table.cell(0, i).text = h
            fail_table.cell(0, i).paragraphs[0].runs[0].bold = True
        for tc in context["top_failures"]:
            row = fail_table.add_row()
            row.cells[0].text = tc.name
            row.cells[1].text = tc.category or "—"
            row.cells[2].text = tc.status.upper()
            row.cells[3].text = (tc.failure_message or "—")[:150]

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"ValReport v1.0 — UST Engineering / Intel Platform Validation\n"
        f"{'Approved by: ' + context['approved_by'] if context.get('approved_by') else 'Status: Draft — Pending Review'}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    filename = f"report_{context['build_id']}_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    return output_path